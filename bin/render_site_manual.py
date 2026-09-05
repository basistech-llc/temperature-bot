#!/usr/bin/env python3
"""Render ``doc/site-manual.md`` to ``doc/site-manual.docx``.

The Markdown file is the source of truth. This script exists so the Word
version can be thrown away and rebuilt rather than hand-edited, which is the
only way to keep the two from drifting apart.

Run it through the Makefile::

    make site-manual-docx

The subset of Markdown understood here is the subset the site manual uses:
``##`` and ``###`` headings, paragraphs, ``-`` bullets with one level of
nesting, fenced code blocks, pipe tables, and inline ``**bold**``, ``*italic*``
and ``code`` spans. Anything else passes through as plain text.

Table column widths are computed rather than shared out evenly. Word wraps
without hyphenating, so a column narrower than its longest unbreakable token
splits words down the middle -- ``AeotecMulti/Sensor6``. Each column therefore
claims at least the width of its longest token before the remaining space is
allocated in proportion to how much text the column actually holds.
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph

ROOT_DIR = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT_DIR / "doc" / "site-manual.md"
DEFAULT_OUTPUT = ROOT_DIR / "doc" / "site-manual.docx"

# Palette. Deep slate for structure, near-black for prose, grey for the rest.
INK = RGBColor(0x25, 0x30, 0x3B)
SLATE = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x6B, 0x7A, 0x8A)
RULE = "C9D2DB"
BAND = "EEF2F6"
ZEBRA = "F8FAFB"
CODEBG = "F1F3F5"

SERIF = "Georgia"
SANS = "Trebuchet MS"
MONO = "Consolas"

# US Letter in EMU. python-docx works in EMU; 914400 EMU == 1 inch.
PAGE_WIDTH = Emu(int(8.5 * 914400))
PAGE_HEIGHT = Emu(int(11 * 914400))
MARGIN = Emu(int(0.8 * 914400))
CONTENT_WIDTH = Emu(PAGE_WIDTH - 2 * MARGIN)

# Rough advance width of one 8pt Georgia character, used only to keep a column
# wide enough for its longest word. Approximate on purpose: it is a floor.
CHAR_EMU = 56000
CELL_PADDING_EMU = 137000
MAX_COLUMN_CHARS = 34

INLINE_RE = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+`|\*[^*]+?\*)")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
SEPARATOR_CELL_RE = re.compile(r"^:?-{2,}:?$")


@dataclass(frozen=True)
class Block:
    """One parsed piece of Markdown, ready to be written out."""

    kind: str
    text: str = ""
    lines: tuple[str, ...] = ()
    rows: tuple[tuple[str, ...], ...] = ()
    depth: int = 0


@dataclass(frozen=True)
class Style:
    """The formatting a run inherits from its surroundings."""

    font: str = SERIF
    size: float = 10.5
    color: RGBColor = INK
    bold: bool = False
    italic: bool = False


# WordprocessingML validates child elements as an ordered sequence, so a
# correct element in the wrong position is a schema error. These are the
# orders for the containers this script writes into by hand; python-docx's
# own ``get_or_add_*`` helpers already place the elements they know about.
_CHILD_ORDER = {
    "w:pPr": (
        "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
        "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
        "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
        "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
        "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
        "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
        "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
        "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
    ),
    "w:tcPr": (
        "w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge", "w:tcBorders",
        "w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText",
        "w:vAlign", "w:hideMark",
    ),
    "w:rPr": (
        "w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
        "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
        "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
        "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern", "w:position",
        "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd",
        "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang",
    ),
    "w:tcMar": ("w:top", "w:start", "w:left", "w:bottom", "w:end", "w:right"),
}


def _insert_ordered(parent, child) -> None:
    """Insert ``child`` at the position the schema requires, replacing any twin."""
    tag = parent.tag.split("}")[-1]
    order = _CHILD_ORDER.get(f"w:{tag}")
    for existing in parent.findall(child.tag):
        parent.remove(existing)
    if order is None:
        parent.append(child)
        return
    name = child.tag.split("}")[-1]
    rank = order.index(f"w:{name}") if f"w:{name}" in order else len(order)
    for sibling in parent:
        sibling_name = f"w:{sibling.tag.split('}')[-1]}"
        sibling_rank = order.index(sibling_name) if sibling_name in order else len(order)
        if sibling_rank > rank:
            sibling.addprevious(child)
            return
    parent.append(child)


def _set_shading(element, fill: str) -> None:
    """Apply a solid background fill to a run, paragraph, or table cell."""
    shd = element.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _insert_ordered(element, shd)


def _set_borders(element, **edges: tuple[int, str] | None) -> None:
    """Set paragraph or cell borders. Each edge is ``(eighth_points, hex)``."""
    tag = "w:pBdr" if element.tag == qn("w:pPr") else "w:tcBorders"
    holder = element.makeelement(qn(tag), {})
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge)
        border = holder.makeelement(qn(f"w:{edge}"), {})
        if spec is None:
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:color"), "auto")
        else:
            size, color = spec
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(size))
            border.set(qn("w:space"), "1")
            border.set(qn("w:color"), color)
        holder.append(border)
    _insert_ordered(element, holder)


def _set_spacing(paragraph: Paragraph, before: int = 0, after: int = 0, line: int = 0) -> None:
    """Set paragraph spacing in twentieths of a point."""
    spacing = paragraph.paragraph_format.element.get_or_add_pPr().get_or_add_spacing()
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def _keep_with_next(paragraph: Paragraph) -> None:
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    _insert_ordered(pPr, pPr.makeelement(qn("w:keepNext"), {}))


def _add_runs(paragraph: Paragraph, text: str, style: Style) -> None:
    """Append ``text`` to ``paragraph``, honouring nested inline markup.

    Recursion matters: ``**`611` and `612` are child devices**`` has code
    spans inside a bold span, and a flat scan would print the backticks.
    """
    position = 0
    for match in INLINE_RE.finditer(text):
        _add_plain(paragraph, text[position : match.start()], style)
        token = match.group(0)
        if token.startswith("**"):
            nested = Style(style.font, style.size, style.color, True, style.italic)
            _add_runs(paragraph, token[2:-2], nested)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _apply(run, Style(MONO, style.size - 1.5, style.color, style.bold, style.italic))
            _set_shading(run.element.get_or_add_rPr(), CODEBG)
        else:
            nested = Style(style.font, style.size, style.color, style.bold, True)
            _add_runs(paragraph, token[1:-1], nested)
        position = match.end()
    _add_plain(paragraph, text[position:], style)


def _add_plain(paragraph: Paragraph, text: str, style: Style) -> None:
    if text:
        _apply(paragraph.add_run(text), style)


def _apply(run, style: Style) -> None:
    run.font.name = style.font
    run.font.size = Pt(style.size)
    run.font.color.rgb = style.color
    run.font.bold = style.bold
    run.font.italic = style.italic
    # Word consults a separate attribute for non-Latin runs; without it the
    # font silently reverts for anything outside the Latin range.
    rFonts = run.element.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), style.font)


def _plain_text(cell: str) -> str:
    return cell.replace("`", "").replace("*", "")


def _column_widths(rows: list[list[str]]) -> list[int]:
    """Allocate the content width across columns, in EMU."""
    columns = len(rows[0])
    longest_token = []
    wanted = []
    for index in range(columns):
        cells = [_plain_text(row[index]) if index < len(row) else "" for row in rows]
        longest_token.append(max((len(word) for cell in cells for word in cell.split()), default=1))
        wanted.append(min(MAX_COLUMN_CHARS, max((len(cell) for cell in cells), default=1)))

    floors = [chars * CHAR_EMU + CELL_PADDING_EMU for chars in longest_token]
    total_wanted = sum(wanted) or 1
    widths = [
        max(floors[i], int(CONTENT_WIDTH * wanted[i] / total_wanted)) for i in range(columns)
    ]

    # Claiming the floor can overshoot the page. Take the excess back from
    # whichever columns still sit above their own floor.
    overflow = sum(widths) - CONTENT_WIDTH
    while overflow > 0:
        slack = [max(0, widths[i] - floors[i]) for i in range(columns)]
        available = sum(slack)
        if available <= 0:
            break
        widths = [widths[i] - int(overflow * slack[i] / available) for i in range(columns)]
        remaining = sum(widths) - CONTENT_WIDTH
        if remaining >= overflow:
            break
        overflow = remaining

    widths[-1] += CONTENT_WIDTH - sum(widths)
    return widths


def _fill_cell(cell: _Cell, text: str, width: int, *, header: bool, fill: str | None) -> None:
    cell.width = Emu(width)
    paragraph = cell.paragraphs[0]
    style = Style(SANS if header else SERIF, 8, SLATE if header else INK, bold=header)
    _add_runs(paragraph, text, style)
    _set_spacing(paragraph, line=250)

    tcPr = cell._tc.get_or_add_tcPr()
    if fill:
        _set_shading(tcPr, fill)
    edge = (4, "1F3A5F") if header else (2, RULE)
    _set_borders(tcPr, top=edge, bottom=edge, left=None, right=None)

    margins = tcPr.makeelement(qn("w:tcMar"), {})
    for side, value in (("top", 40), ("start", 60), ("bottom", 40), ("end", 60)):
        node = margins.makeelement(qn(f"w:{side}"), {})
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    _insert_ordered(tcPr, margins)


def _add_table(document: DocumentObject, rows: list[list[str]]) -> None:
    widths = _column_widths(rows)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    grid = table._tbl.find(qn("w:tblGrid"))
    for column, width in zip(grid.findall(qn("w:gridCol")), widths):
        column.set(qn("w:w"), str(int(Emu(width).twips)))

    for row_index, row in enumerate(rows):
        header = row_index == 0
        if header:
            # Repeat the header when a long inventory spills onto a new page.
            trPr = table.rows[0]._tr.get_or_add_trPr()
            _insert_ordered(trPr, trPr.makeelement(qn("w:tblHeader"), {}))
        fill = BAND if header else (ZEBRA if row_index % 2 == 0 else None)
        for column_index in range(len(rows[0])):
            text = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            _fill_cell(cell, text, widths[column_index], header=header, fill=fill)

    _set_spacing(document.add_paragraph(), after=200)


def _add_page_number(paragraph: Paragraph) -> None:
    """Insert a PAGE field. python-docx has no API for these."""
    run = paragraph.add_run()
    begin = run.element.makeelement(qn("w:fldChar"), {})
    begin.set(qn("w:fldCharType"), "begin")
    instruction = run.element.makeelement(qn("w:instrText"), {})
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = run.element.makeelement(qn("w:fldChar"), {})
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, end):
        run.element.append(node)
    _apply(run, Style(SANS, 8.5, MUTED))


def _parse(lines: list[str]) -> list[Block]:
    """Turn Markdown lines into a flat list of blocks."""
    blocks: list[Block] = []
    index = 0
    while index < len(lines):
        line = lines[index]

        if line.startswith("# "):
            index += 1
        elif line.startswith("## "):
            blocks.append(Block("h1", text=line[3:].strip()))
            index += 1
        elif line.startswith("### "):
            blocks.append(Block("h2", text=line[4:].strip()))
            index += 1
        elif line.startswith("```"):
            index += 1
            code: list[str] = []
            while index < len(lines) and not lines[index].startswith("```"):
                code.append(lines[index])
                index += 1
            index += 1
            blocks.append(Block("code", lines=tuple(code)))
        elif TABLE_ROW_RE.match(line):
            rows: list[list[str]] = []
            while index < len(lines) and TABLE_ROW_RE.match(lines[index]):
                cells = [c.strip() for c in lines[index].strip().strip("|").split("|")]
                if not all(SEPARATOR_CELL_RE.match(c) for c in cells):
                    rows.append(cells)
                index += 1
            blocks.append(Block("table", rows=tuple(tuple(r) for r in rows)))
        elif re.match(r"^\s*- ", line):
            while index < len(lines) and re.match(r"^\s*- |^\s{3,}\S", lines[index]):
                depth = 1 if re.match(r"^\s{2,}- ", lines[index]) else 0
                text = re.sub(r"^\s*- ", "", lines[index]).strip()
                index += 1
                while (
                    index < len(lines)
                    and re.match(r"^\s{2,}\S", lines[index])
                    and not re.match(r"^\s*- ", lines[index])
                ):
                    text += " " + lines[index].strip()
                    index += 1
                blocks.append(Block("bullet", text=text, depth=depth))
        elif not line.strip():
            index += 1
        else:
            paragraph = [line.strip()]
            index += 1
            while (
                index < len(lines)
                and lines[index].strip()
                and not re.match(r"^(#|\||\s*- |```)", lines[index])
            ):
                paragraph.append(lines[index].strip())
                index += 1
            blocks.append(Block("para", text=" ".join(paragraph)))
    return blocks


def _write_front_matter(
    document: DocumentObject, lede: list[Block], toc: list[str]
) -> None:
    _set_spacing(document.add_paragraph(), after=2600)

    eyebrow = document.add_paragraph()
    run = eyebrow.add_run("TEMPERATURE BOT")
    _apply(run, Style(SANS, 10, MUTED))
    rPr = run.element.get_or_add_rPr()
    spacing = rPr.makeelement(qn("w:spacing"), {})
    spacing.set(qn("w:val"), "80")
    _insert_ordered(rPr, spacing)
    _set_spacing(eyebrow, after=120)

    title = document.add_paragraph()
    _apply(title.add_run("Site Manual"), Style(SANS, 34, SLATE, bold=True))
    _set_spacing(title, after=60)

    divider = document.add_paragraph()
    _set_borders(
        divider.paragraph_format.element.get_or_add_pPr(),
        bottom=(12, "1F3A5F"),
    )

    subtitle = document.add_paragraph()
    _apply(
        subtitle.add_run("The Hubitat installation, hub by hub and device by device"),
        Style(SERIF, 13, MUTED, italic=True),
    )
    _set_spacing(subtitle, before=200, after=1400)

    for text, style in (
        ("Surveyed 20 August 2026", Style(SANS, 10, INK)),
        ("Hubs 10.2.3.51 - 10.2.3.54, Somerville", Style(SANS, 10, MUTED)),
    ):
        paragraph = document.add_paragraph()
        _apply(paragraph.add_run(text), style)
        _set_spacing(paragraph, after=60)

    caveat = document.add_paragraph()
    _apply(
        caveat.add_run("A snapshot. Regenerate from the hubs before relying on any number."),
        Style(SERIF, 9.5, MUTED, italic=True),
    )
    _set_spacing(caveat, before=240)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    heading = document.add_paragraph()
    _apply(heading.add_run("Contents"), Style(SANS, 15, SLATE, bold=True))
    _set_spacing(heading, after=60)
    _set_borders(heading.paragraph_format.element.get_or_add_pPr(), bottom=(6, RULE))
    _set_spacing(document.add_paragraph(), after=180)

    for entry in toc:
        paragraph = document.add_paragraph()
        _apply(paragraph.add_run(entry), Style(SERIF, 10.5, INK))
        _set_spacing(paragraph, after=130)
        paragraph.paragraph_format.left_indent = Emu(127000)

    _set_spacing(document.add_paragraph(), after=200)
    _write_blocks(document, lede, first_heading_breaks=False)


def _write_blocks(
    document: DocumentObject, blocks: list[Block], *, first_heading_breaks: bool
) -> None:
    seen_heading = False
    for block in blocks:
        if block.kind == "h1":
            paragraph = document.add_paragraph()
            if seen_heading or first_heading_breaks:
                paragraph.add_run().add_break(WD_BREAK.PAGE)
            seen_heading = True
            _apply(paragraph.add_run(block.text), Style(SANS, 15, SLATE, bold=True))
            _set_spacing(paragraph, before=0 if not seen_heading else 60, after=60)
            _set_borders(paragraph.paragraph_format.element.get_or_add_pPr(), bottom=(6, RULE))
            _keep_with_next(paragraph)
        elif block.kind == "h2":
            paragraph = document.add_paragraph()
            _apply(paragraph.add_run(block.text), Style(SANS, 11, SLATE, bold=True))
            _set_spacing(paragraph, before=300, after=100)
            _keep_with_next(paragraph)
        elif block.kind == "para":
            paragraph = document.add_paragraph()
            _add_runs(paragraph, block.text, Style())
            _set_spacing(paragraph, after=160, line=300)
        elif block.kind == "bullet":
            depth = block.depth
            paragraph = document.add_paragraph()
            marker = "\u2013 " if depth == 0 else "\u00b7 "
            _apply(paragraph.add_run(marker), Style(SERIF, 10.5, SLATE if depth == 0 else MUTED))
            _add_runs(paragraph, block.text, Style())
            _set_spacing(paragraph, after=90, line=290)
            # 380 twips in, 220 back out, so wrapped lines clear the marker.
            paragraph.paragraph_format.left_indent = Emu(241300 if depth == 0 else 482600)
            paragraph.paragraph_format.first_line_indent = Emu(-139700)
        elif block.kind == "code":
            for offset, text in enumerate(block.lines):
                paragraph = document.add_paragraph()
                _apply(paragraph.add_run(text or " "), Style(MONO, 8, INK))
                _set_shading(paragraph.paragraph_format.element.get_or_add_pPr(), CODEBG)
                _set_spacing(
                    paragraph,
                    before=60 if offset == 0 else 0,
                    after=200 if offset == len(block.lines) - 1 else 0,
                    line=240,
                )
                paragraph.paragraph_format.left_indent = Emu(127000)
                paragraph.paragraph_format.right_indent = Emu(127000)
        elif block.kind == "table":
            _add_table(document, [list(row) for row in block.rows])


def render(source: Path, output: Path) -> None:
    blocks = _parse(source.read_text(encoding="utf-8").split("\n"))
    toc = [block.text for block in blocks if block.kind == "h1"]

    first_heading = next((i for i, block in enumerate(blocks) if block.kind == "h1"), 0)
    lede, body = blocks[:first_heading], blocks[first_heading:]

    document = Document()
    document.core_properties.title = "Site Manual"
    document.core_properties.author = "Temperature Bot"
    document.core_properties.comments = "Generated from doc/site-manual.md; do not hand-edit."

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    for side in ("top", "bottom", "left", "right"):
        setattr(section, f"{side}_margin", MARGIN)

    # A title page with no page number, numbering from the contents onward.
    section.different_first_page_header_footer = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer)

    normal = document.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(0)

    # python-docx ships a settings.xml whose <w:zoom> omits the percent
    # attribute the schema requires. Harmless in practice, but this file is
    # committed, so leave it validating cleanly.
    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    _write_front_matter(document, lede, toc)
    _write_blocks(document, body, first_heading_breaks=True)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args(argv)

    if not args.source.exists():
        print(f"No such file: {args.source}", file=sys.stderr)
        return 1

    render(args.source, args.output)
    print(f"Wrote {args.output.relative_to(ROOT_DIR)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
